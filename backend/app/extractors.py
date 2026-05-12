from __future__ import annotations

import csv
import io
import tempfile
from pathlib import Path
from typing import Any


def extract_content(file_name: str, extension: str, content: bytes) -> dict[str, Any]:
    extension = extension.lower().lstrip(".")
    if extension == "pdf":
        return extract_pdf(file_name, content)
    if extension in {"xlsx", "xlsm", "xls"}:
        return extract_excel(file_name, content)
    if extension == "docx":
        return extract_docx(file_name, content)
    if extension in {"txt", "md", "markdown"}:
        return extract_text_file(file_name, extension, content)
    raise ValueError(f"UNSUPPORTED_FILE_TYPE: 不支持的文件格式 {extension}")


def extract_text_file(file_name: str, extension: str, content: bytes) -> dict[str, Any]:
    warnings: list[str] = []
    try:
        from charset_normalizer import from_bytes

        result = from_bytes(content).best()
        text = str(result) if result else content.decode("utf-8", errors="replace")
        encoding = result.encoding if result else "utf-8"
    except Exception:
        text = content.decode("utf-8", errors="replace")
        encoding = "utf-8/replace"
        warnings.append("编码自动识别失败，已使用 UTF-8 容错解码。")
    cleaned = _clean_text(text)
    return {
        "text": f"# 文件：{file_name}\n类型：{extension.upper()} 文本\n提取方式：文本解码（{encoding}）\n\n{cleaned}".strip(),
        "engine": "charset-normalizer",
        "warnings": warnings,
    }


def extract_docx(file_name: str, content: bytes) -> dict[str, Any]:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("DOCX_ENGINE_MISSING: 缺少 python-docx 依赖，无法解析 Word 文档。") from exc

    doc = Document(io.BytesIO(content))
    parts = [f"# 文件：{file_name}", "类型：Word DOCX", "提取方式：python-docx", ""]
    for paragraph in doc.paragraphs:
        text = _clean_text(paragraph.text)
        if text:
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            if "heading" in style or "标题" in style:
                parts.append(f"## {text}")
            else:
                parts.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"\n## 表格 {table_index}")
        rows = [[_clean_text(cell.text) for cell in row.cells] for row in table.rows]
        parts.append(_rows_to_markdown_or_tsv(rows))
    text = "\n\n".join(part for part in parts if part is not None).strip()
    return {"text": text, "engine": "python-docx", "warnings": []}


def extract_excel(file_name: str, content: bytes) -> dict[str, Any]:
    if file_name.lower().endswith(".xls"):
        return extract_excel_with_pandas(file_name, content)
    try:
        from openpyxl import load_workbook
    except Exception as exc:
        raise RuntimeError("EXCEL_ENGINE_MISSING: 缺少 openpyxl 依赖，无法解析 Excel。") from exc

    warnings: list[str] = []
    workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=False)
    parts = [f"# 文件：{file_name}", "类型：Excel", "提取方式：openpyxl（读取单元格显示值）", ""]
    for sheet in workbook.worksheets:
        parts.append(f"## Sheet：{sheet.title}")
        merged_ranges = list(sheet.merged_cells.ranges)
        if merged_ranges:
            warnings.append(f"Sheet「{sheet.title}」包含合并单元格，已将左上角值补齐到合并区域。")
        values_by_coord = {}
        for merged_range in merged_ranges:
            value = sheet.cell(merged_range.min_row, merged_range.min_col).value
            for row in range(merged_range.min_row, merged_range.max_row + 1):
                for col in range(merged_range.min_col, merged_range.max_col + 1):
                    values_by_coord[(row, col)] = value
        rows: list[list[str]] = []
        for row in sheet.iter_rows():
            values = []
            for cell in row:
                value = values_by_coord.get((cell.row, cell.column), cell.value)
                values.append(_cell_to_text(value))
            if any(value for value in values):
                rows.append(values)
        if not rows:
            parts.append("（空 Sheet）")
            continue
        parts.append(_rows_to_markdown_or_tsv(rows))
    return {"text": "\n\n".join(parts).strip(), "engine": "openpyxl", "warnings": warnings}


def extract_excel_with_pandas(file_name: str, content: bytes) -> dict[str, Any]:
    try:
        import pandas as pd
    except Exception as exc:
        raise RuntimeError("EXCEL_ENGINE_MISSING: 缺少 pandas/xlrd 依赖，无法解析 xls。") from exc
    warnings = ["该文件为旧版 xls，已使用 pandas/xlrd 提取单元格显示值。"]
    sheets = pd.read_excel(io.BytesIO(content), sheet_name=None, header=None, dtype=str)
    parts = [f"# 文件：{file_name}", "类型：Excel", "提取方式：pandas/xlrd（读取单元格显示值）", ""]
    for sheet_name, frame in sheets.items():
        parts.append(f"## Sheet：{sheet_name}")
        rows = frame.fillna("").astype(str).values.tolist()
        parts.append(_rows_to_markdown_or_tsv(rows))
    return {"text": "\n\n".join(parts).strip(), "engine": "pandas/xlrd", "warnings": warnings}


def extract_pdf(file_name: str, content: bytes) -> dict[str, Any]:
    warnings: list[str] = []
    page_texts: list[str] = []
    table_texts: list[str] = []
    engine = "PyMuPDF + pdfplumber"

    try:
        import fitz
    except Exception as exc:
        raise RuntimeError("PDF_ENGINE_MISSING: 缺少 PyMuPDF 依赖，无法解析 PDF。") from exc

    with fitz.open(stream=content, filetype="pdf") as doc:
        for page_index, page in enumerate(doc, start=1):
            text = _clean_text(page.get_text("text") or "")
            if text:
                page_texts.append(f"## 第 {page_index} 页\n{text}")

    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables() or []
                for table_index, table in enumerate(tables, start=1):
                    table_texts.append(f"### 第 {page_index} 页 表格 {table_index}\n{_rows_to_markdown_or_tsv(table)}")
    except Exception as exc:
        warnings.append(f"pdfplumber 表格提取失败：{exc}")

    text_char_count = sum(len(item) for item in page_texts)
    scanned_like = text_char_count < 50
    if scanned_like:
        warnings.append("该 PDF 可能为扫描件，已尝试使用 PaddleOCR 进行 OCR。")
        ocr_result = _extract_pdf_with_paddleocr(content)
        if ocr_result["text"]:
            engine = "PaddleOCR"
            page_texts = [ocr_result["text"]]
            warnings.extend(ocr_result["warnings"])
        else:
            warnings.extend(ocr_result["warnings"])

    header = [
        f"# 文件：{file_name}",
        "类型：PDF",
        f"提取方式：{engine}",
    ]
    if scanned_like:
        header.append("识别提示：该文件疑似扫描件，OCR 结果可能存在错别字或漏识别，请以原件为准。")
    body = "\n\n".join(page_texts + table_texts).strip()
    if not body:
        body = "（未提取到可用文本）"
    return {"text": "\n".join(header).strip() + "\n\n" + body, "engine": engine, "warnings": warnings}


def _extract_pdf_with_paddleocr(content: bytes) -> dict[str, Any]:
    warnings: list[str] = []
    try:
        import fitz
        from paddleocr import PaddleOCR
    except Exception as exc:
        return {"text": "", "warnings": [f"PaddleOCR 未安装或不可用，无法识别扫描 PDF：{exc}"]}

    texts: list[str] = []
    with tempfile.TemporaryDirectory() as tmpdir:
        ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page_index, page in enumerate(doc, start=1):
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image_path = Path(tmpdir) / f"page_{page_index}.png"
                pixmap.save(str(image_path))
                result = ocr.ocr(str(image_path), cls=True)
                lines: list[str] = []
                for block in result or []:
                    for item in block or []:
                        if len(item) >= 2 and isinstance(item[1], (list, tuple)):
                            lines.append(str(item[1][0]))
                if lines:
                    texts.append(f"## 第 {page_index} 页 OCR文本\n" + "\n".join(lines))
    return {"text": "\n\n".join(texts), "warnings": warnings}


def _rows_to_markdown_or_tsv(rows: list[list[Any]]) -> str:
    normalized = [[_cell_to_text(cell) for cell in row] for row in rows if row is not None]
    normalized = [row for row in normalized if any(cell for cell in row)]
    if not normalized:
        return "（空表格）"
    max_cols = max(len(row) for row in normalized)
    normalized = [row + [""] * (max_cols - len(row)) for row in normalized]
    if len(normalized) <= 40 and max_cols <= 12:
        header = normalized[0]
        body = normalized[1:] or [[""] * max_cols]
        lines = [
            "| 行号 | " + " | ".join(_escape_md(cell or f"列{index + 1}") for index, cell in enumerate(header)) + " |",
            "|---:|" + "|".join("---" for _ in range(max_cols)) + "|",
        ]
        for row_index, row in enumerate(body, start=2):
            lines.append(f"| {row_index} | " + " | ".join(_escape_md(cell) for cell in row) + " |")
        return "\n".join(lines)
    output = io.StringIO()
    writer = csv.writer(output, delimiter="\t", lineterminator="\n")
    writer.writerow(["row_id", *[cell or f"col_{index + 1}" for index, cell in enumerate(normalized[0])]])
    for row_index, row in enumerate(normalized[1:], start=2):
        writer.writerow([row_index, *row])
    return output.getvalue().strip()


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    return _clean_text(str(value))


def _clean_text(value: str) -> str:
    return "\n".join(line.strip() for line in value.replace("\x00", "").splitlines()).strip()


def _escape_md(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>")
